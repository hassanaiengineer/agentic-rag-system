from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import fitz
import numpy as np
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from PIL import Image, ImageFilter, ImageOps
from pdf2image import convert_from_path
import pytesseract
import PyPDF2


@dataclass
class ExtractedBlock:
    text: str
    bbox: list[float]


@dataclass
class ExtractedPage:
    page: int
    blocks: list[ExtractedBlock]


@dataclass
class StructuredDocument:
    text: str
    pages: list[ExtractedPage]
    processing_path: str


class FileProcessor:
    @staticmethod
    def pre_process_pdf(file_path: Path) -> None:
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                if reader.is_encrypted:
                    raise ValueError(f"File {file_path.name} is encrypted and cannot be processed.")
        except Exception as e:
            if "encrypted" in str(e):
                raise
            # If PyPDF2 fails for other reasons, we still let PyMuPDF try.
            pass


class IngestionService:
    def __init__(self, min_page_chars: int = 100, density_threshold: float = 0.0005):
        self.min_page_chars = min_page_chars
        self.density_threshold = density_threshold

    @staticmethod
    def normalize_text(text: str) -> str:
        cleaned = text.replace("\x00", " ")
        cleaned = re.sub(r"[ \t]+", " ", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        cleaned = re.sub(r"[^\S\n]+", " ", cleaned)
        return cleaned.strip()

    def extract_structured(self, file_path: Path) -> StructuredDocument:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            FileProcessor.pre_process_pdf(file_path)
            return self._extract_pdf_hybrid(file_path)
        if suffix == ".docx":
            return self._extract_docx_structured(file_path)
        raise ValueError(f"Unsupported file type: {suffix}")

    def extract_text(self, file_path: Path) -> str:
        return self.extract_structured(file_path).text

    def _extract_pdf_hybrid(self, file_path: Path) -> StructuredDocument:
        digital = self._extract_pdf_digital(file_path)
        if self._is_low_quality_pdf(digital):
            return self._extract_pdf_ocr(file_path)
        return digital

    def _extract_pdf_digital(self, file_path: Path) -> StructuredDocument:
        pages: list[ExtractedPage] = []
        with fitz.open(file_path) as doc:
            for i, page in enumerate(doc, start=1):
                raw = page.get_text("blocks")
                blocks = self._sort_blocks_reading_order(raw)
                pages.append(ExtractedPage(page=i, blocks=blocks))

        full_text = self._assemble_document_text(pages)
        return StructuredDocument(text=full_text, pages=pages, processing_path="digital")

    def _extract_pdf_ocr(self, file_path: Path) -> StructuredDocument:
        images = convert_from_path(str(file_path), dpi=250)
        pages: list[ExtractedPage] = []

        for i, image in enumerate(images, start=1):
            prepped = self._preprocess_image(image)
            data = pytesseract.image_to_data(prepped, output_type=pytesseract.Output.DICT)
            blocks = self._ocr_data_to_blocks(data)
            pages.append(ExtractedPage(page=i, blocks=blocks))

        full_text = self._assemble_document_text(pages)
        return StructuredDocument(text=full_text, pages=pages, processing_path="ocr")

    def _extract_docx_structured(self, file_path: Path) -> StructuredDocument:
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            raise ValueError(f"File {file_path.name} is corrupted or not a valid DOCX file.")
        except Exception as e:
            raise ValueError(f"Error parsing DOCX file {file_path.name}: {e}")

        blocks: list[ExtractedBlock] = []
        for para in doc.paragraphs:
            txt = self.normalize_text(para.text)
            if txt:
                blocks.append(ExtractedBlock(text=txt, bbox=[0.0, 0.0, 0.0, 0.0]))

        page = ExtractedPage(page=1, blocks=blocks)
        text = self._assemble_document_text([page])
        return StructuredDocument(text=text, pages=[page], processing_path="digital")

    def _is_low_quality_pdf(self, doc: StructuredDocument) -> bool:
        if not doc.pages:
            return True

        total_chars = len(doc.text)
        avg_chars = total_chars / max(1, len(doc.pages))
        if avg_chars < self.min_page_chars:
            return True

        block_area = 0.0
        for page in doc.pages:
            for block in page.blocks:
                x1, y1, x2, y2 = block.bbox
                block_area += max(0.0, (x2 - x1) * (y2 - y1))
        density = total_chars / max(1.0, block_area)
        return density < self.density_threshold

    def _sort_blocks_reading_order(self, raw_blocks: list[tuple]) -> list[ExtractedBlock]:
        items: list[ExtractedBlock] = []
        for b in raw_blocks:
            x1, y1, x2, y2, text = b[:5]
            normalized = self.normalize_text(text)
            if not normalized:
                continue
            items.append(ExtractedBlock(text=normalized, bbox=[float(x1), float(y1), float(x2), float(y2)]))

        tolerance = 6.0
        items.sort(key=lambda blk: (round(blk.bbox[1] / tolerance), blk.bbox[0]))
        return items

    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        gray = ImageOps.grayscale(image)
        width, height = gray.size
        if width < 1400:
            scale = 1400 / max(1, width)
            gray = gray.resize((int(width * scale), int(height * scale)))

        arr = np.array(gray)
        threshold = np.where(arr > 180, 255, 0).astype(np.uint8)
        binarized = Image.fromarray(threshold)
        sharpened = binarized.filter(ImageFilter.SHARPEN)
        return sharpened

    def _ocr_data_to_blocks(self, data: dict) -> list[ExtractedBlock]:
        blocks: list[ExtractedBlock] = []
        n = len(data.get("text", []))
        line_map: dict[tuple[int, int, int], list[int]] = {}

        for i in range(n):
            txt = self.normalize_text(str(data["text"][i]))
            conf = int(data.get("conf", ["-1"])[i]) if str(data.get("conf", ["-1"])[i]).isdigit() else -1
            if not txt or conf < 30:
                continue
            key = (int(data["block_num"][i]), int(data["par_num"][i]), int(data["line_num"][i]))
            line_map.setdefault(key, []).append(i)

        for _, idxs in sorted(line_map.items(), key=lambda kv: kv[0]):
            texts = [str(data["text"][i]).strip() for i in idxs]
            text = self.normalize_text(" ".join(texts))
            if not text:
                continue
            x1 = min(int(data["left"][i]) for i in idxs)
            y1 = min(int(data["top"][i]) for i in idxs)
            x2 = max(int(data["left"][i]) + int(data["width"][i]) for i in idxs)
            y2 = max(int(data["top"][i]) + int(data["height"][i]) for i in idxs)
            blocks.append(ExtractedBlock(text=text, bbox=[float(x1), float(y1), float(x2), float(y2)]))

        blocks.sort(key=lambda blk: (blk.bbox[1], blk.bbox[0]))
        return blocks

    def _assemble_document_text(self, pages: list[ExtractedPage]) -> str:
        joined_pages: list[str] = []
        for page in pages:
            lines = [blk.text for blk in page.blocks if blk.text]
            page_text = self.normalize_text("\n".join(lines))
            if page_text:
                joined_pages.append(page_text)
        return self.normalize_text("\n\n".join(joined_pages))
